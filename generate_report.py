"""Generate FitTrack Pro Technical Report as Word document."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Page margins (approx 2-page layout)
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title
title = doc.add_heading('FitTrack Pro — Technical Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('Project: Personal Fitness & Health Tracker Website\n')
run.bold = True
meta.add_run('Technologies: HTML5, CSS3, JavaScript (Vanilla)\n')
meta.add_run('Date: June 2026')

doc.add_paragraph()

# Section 1
doc.add_heading('1. Problem-Solving Approach', level=1)

doc.add_heading('1.1 Problem Definition', level=2)
doc.add_paragraph(
    'The goal was to build an all-in-one fitness platform where users can track workouts, '
    'body metrics (BMI), hydration, and nutrition without installing an app or creating an account. '
    'The solution had to be lightweight, work offline after the first load, feel professional, '
    'and remain usable on phones, tablets, and desktops.'
)

doc.add_heading('1.2 Requirements Analysis', level=2)
doc.add_paragraph('Before implementation, the project was broken into clear functional and non-functional requirements:')

req_table = doc.add_table(rows=5, cols=2)
req_table.style = 'Table Grid'
req_data = [
    ('Category', 'Requirements'),
    ('Functional', 'Workout logging, BMI calculation, water tracking (8-glass goal), calorie goal with meal logging, contact form'),
    ('UI/UX', 'Red/black fitness theme, hero section, carousels, responsive navigation, visual feedback on actions'),
    ('Technical', 'No backend dependency, persistent data, semantic HTML, cross-browser compatibility'),
    ('Quality', 'Accessibility for screen readers, input validation, protection against XSS in dynamic content'),
]
for i, (a, b) in enumerate(req_data):
    req_table.rows[i].cells[0].text = a
    req_table.rows[i].cells[1].text = b
    if i == 0:
        for cell in req_table.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

doc.add_paragraph()

doc.add_heading('1.3 Architecture & Design Decisions', level=2)
decisions = [
    ('Multi-page static site', 'Six HTML pages share one CSS file and one JavaScript file. This keeps deployment simple while separating concerns by feature page.'),
    ('Client-side persistence with localStorage', 'A server database was avoided to eliminate hosting complexity and keep user data private. JSON stores workout arrays, meal lists, and daily counters. Date keys ensure water and calorie logs reset automatically each day.'),
    ('Modular JavaScript initialization', 'script.js uses a DOMContentLoaded entry point calling focused init functions. Each checks for required DOM elements before binding events, so the same script loads on every page without errors.'),
    ('Progressive UI enhancement', 'The homepage includes hero carousel, programs carousel with touch swipe, Intersection Observer scroll-reveal animations, and animated stat counters.'),
]
for heading, text in decisions:
    p = doc.add_paragraph()
    p.add_run(heading + '. ').bold = True
    p.add_run(text)

doc.add_heading('1.4 Development Workflow', level=2)
workflow = [
    'Core trackers first — Forms, validation, localStorage CRUD, and table/list rendering.',
    'Visual redesign — Red/black theme, typography (Bebas Neue + Inter), hero section, and carousels.',
    'Reliability fixes — Broken external URLs replaced with verified links and local images.',
    'Polish — SVG icons, contact link styling, tel: links, and responsive footer.',
]
for i, item in enumerate(workflow, 1):
    doc.add_paragraph(f'{i}. {item}', style='List Number')

doc.add_heading('1.5 Key Algorithms', level=2)
doc.add_paragraph('BMI Calculation: BMI = weight (kg) / (height in m)² with category mapping and visual scale indicator.')
doc.add_paragraph('Daily reset logic: Current date compared to stored date; counters reset at midnight automatically.')
doc.add_paragraph('XSS-safe rendering: User input passes through escapeHTML() before DOM insertion.')

doc.add_page_break()

# Section 2
doc.add_heading('2. Accessibility & Security Features', level=1)

doc.add_heading('2.1 Accessibility (WCAG-Oriented Practices)', level=2)
a11y_table = doc.add_table(rows=13, cols=2)
a11y_table.style = 'Table Grid'
a11y_data = [
    ('Feature', 'Implementation'),
    ('Skip navigation', 'Skip link allows keyboard users to jump to main content'),
    ('Semantic HTML', 'header, nav, main, section, article, footer used throughout'),
    ('ARIA labels', 'Navigation, carousels, and tables have descriptive aria-label attributes'),
    ('Current page', 'aria-current="page" on active navigation links'),
    ('Mobile menu', 'Toggle uses aria-expanded updated on open/close'),
    ('Live regions', 'aria-live="polite" on dynamic workout, water, BMI, and meal updates'),
    ('Progress bars', 'role="progressbar" with aria-valuenow, min, and max values'),
    ('Form accessibility', 'Labels, required fields, aria-required, and autocomplete attributes'),
    ('Decorative content', 'Visual-only elements marked aria-hidden="true"'),
    ('Focus states', 'Visible focus styles on buttons, links, and inputs'),
    ('Color contrast', 'Light text on dark backgrounds with red accent colors'),
    ('Responsive text', 'clamp() for fluid heading sizes across viewports'),
]
for i, (a, b) in enumerate(a11y_data):
    a11y_table.rows[i].cells[0].text = a
    a11y_table.rows[i].cells[1].text = b
    if i == 0:
        for cell in a11y_table.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

doc.add_paragraph()

doc.add_heading('2.2 Security Features', level=2)
sec_table = doc.add_table(rows=7, cols=2)
sec_table.style = 'Table Grid'
sec_data = [
    ('Threat', 'Mitigation'),
    ('Cross-Site Scripting (XSS)', 'escapeHTML() sanitizes user text before rendering in dynamic lists'),
    ('Input validation', 'HTML5 required, type email, min/max; JavaScript rejects invalid submissions'),
    ('No server exposure', 'Contact form is client-side only; no data sent to external APIs'),
    ('localStorage scope', 'Data is origin-bound; no credentials stored'),
    ('External links', 'tel: and mailto: used correctly; social links have aria-label'),
    ('No inline scripts', 'All logic in external script.js; no eval() used'),
]
for i, (a, b) in enumerate(sec_data):
    sec_table.rows[i].cells[0].text = a
    sec_table.rows[i].cells[1].text = b
    if i == 0:
        for cell in sec_table.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

doc.add_paragraph()

doc.add_heading('2.3 Limitations & Future Improvements', level=2)
for item in [
    'Contact form does not send email server-side (demo behavior only).',
    'localStorage has no encryption; data visible in browser DevTools.',
    'No Content Security Policy headers without server configuration.',
    'Future: CSP, HTTPS hosting, and optional user authentication for cloud sync.',
]:
    doc.add_paragraph(item, style='List Bullet')

# Section 3
doc.add_heading('3. Performance Optimization Steps', level=1)

doc.add_heading('3.1 Asset Loading', level=2)
perf_table = doc.add_table(rows=6, cols=2)
perf_table.style = 'Table Grid'
perf_data = [
    ('Optimization', 'Detail'),
    ('Lazy loading images', 'loading="lazy" on below-the-fold images'),
    ('Eager loading for LCP', 'Hero images use loading="eager" for faster first paint'),
    ('Local feature images', 'workout.jpg, bmi.jpg, water.jpg hosted locally'),
    ('Single CSS/JS files', 'One stylesheet and one script reduce HTTP requests'),
    ('Google Fonts', 'Limited weight variants loaded via @import'),
]
for i, (a, b) in enumerate(perf_data):
    perf_table.rows[i].cells[0].text = a
    perf_table.rows[i].cells[1].text = b
    if i == 0:
        for cell in perf_table.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

doc.add_paragraph()

doc.add_heading('3.2 CSS & JavaScript Performance', level=2)
for item in [
    'CSS custom properties for theming without duplication.',
    'GPU-friendly transform and opacity animations instead of layout-triggering properties.',
    'Responsive breakpoints at 1024px, 768px, and 480px.',
    'Passive scroll listeners ({ passive: true }) prevent scroll jank.',
    'Intersection Observer runs animations only when elements enter viewport.',
    'Event delegation on delete buttons reduces listener count.',
    'requestAnimationFrame with cubic easing for stat counters.',
    'Carousel timers clear on user interaction to prevent memory leaks.',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('3.3 Production Recommendations', level=2)
for i, item in enumerate([
    'Minify style.css and script.js',
    'Convert images to WebP with fallbacks',
    'Add preconnect for Google Fonts',
    'Deploy behind HTTPS with gzip/Brotli compression',
    'Run Lighthouse audit and fix remaining warnings',
], 1):
    doc.add_paragraph(f'{i}. {item}', style='List Number')

doc.add_paragraph()

doc.add_heading('Summary', level=1)
doc.add_paragraph(
    'FitTrack Pro demonstrates a structured approach to building a client-side fitness application: '
    'requirements were decomposed by feature, data persistence was solved with localStorage and daily reset logic, '
    'the UI was enhanced with carousels and animations without heavy frameworks, accessibility was built in from '
    'semantic markup and ARIA attributes, XSS risks were mitigated through output encoding, and performance was '
    'improved through lazy loading, passive listeners, Intersection Observer, and CSS transform-based animations. '
    'The result is a deployable, professional fitness website that runs entirely in the browser.'
)

end = doc.add_paragraph('— End of Technical Report —')
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in end.runs:
    r.italic = True

output = r'c:\Users\dell\Desktop\My Pro\fitness-tracker\TECHNICAL_REPORT.docx'
doc.save(output)
print(f'Saved: {output}')
